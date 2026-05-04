# -*- coding: utf-8 -*-
"""
Загрузка и предобработка данных
Поддержка форматов: TXT, DOCX, PDF
"""

import os
import chardet
import pandas as pd
from tqdm import tqdm
import re
from typing import List, Dict, Tuple
import config

# Поддерживаемые расширения файлов
SUPPORTED_EXTENSIONS = ('.txt', '.docx', '.pdf')


def load_txt_file(file_path: str) -> str:
    """Загрузка TXT файла"""
    encodings = ['utf-8', 'windows-1251', 'cp1251', 'utf-16', 'iso-8859-5']
    for encoding in encodings:
        try:
            with open(file_path, 'r', encoding=encoding) as f:
                text = f.read()
                if text and len(text) > config.MIN_TEXT_LENGTH:
                    return text
        except (UnicodeDecodeError, UnicodeError):
            continue
    encoding = detect_encoding(file_path)
    with open(file_path, 'r', encoding=encoding, errors='ignore') as f:
        return f.read()


def load_docx_file(file_path: str) -> str:
    """Загрузка DOCX файла"""
    try:
        from docx import Document
        doc = Document(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        text = '\n'.join(paragraphs)
        for table in doc.tables:
            for row in table.rows:
                text += '\n' + ' '.join(cell.text for cell in row.cells)
        return text.strip()
    except ImportError:
        return ""
    except Exception as e:
        print(f"Ошибка чтения DOCX {file_path}: {e}")
        return ""


def load_pdf_file(file_path: str) -> str:
    """Загрузка PDF файла"""
    try:
        from pypdf import PdfReader
        reader = PdfReader(file_path)
        text_parts = []
        for page in reader.pages[:50]:  # первые 50 страниц
            part = page.extract_text()
            if part:
                text_parts.append(part)
        return '\n'.join(text_parts) if text_parts else ""
    except ImportError:
        return ""
    except Exception as e:
        print(f"Ошибка чтения PDF {file_path}: {e}")
        return ""


def load_file_by_type(file_path: str) -> str:
    """Загрузка файла по типу расширения"""
    ext = os.path.splitext(file_path)[1].lower()
    if ext == '.txt':
        return load_txt_file(file_path)
    elif ext == '.docx':
        return load_docx_file(file_path)
    elif ext == '.pdf':
        return load_pdf_file(file_path)
    return ""


def detect_encoding(file_path: str) -> str:
    """Определяет кодировку файла"""
    with open(file_path, 'rb') as f:
        result = chardet.detect(f.read())
        return result['encoding']


def load_order_file(file_path: str) -> str:
    """Загружает один файл приказа (TXT, DOCX, PDF)"""
    try:
        ext = os.path.splitext(file_path)[1].lower()
        if ext in SUPPORTED_EXTENSIONS:
            return load_file_by_type(file_path)
        return ""
    except Exception as e:
        print(f"Ошибка при чтении файла {file_path}: {e}")
        return ""


def extract_first_n_paragraphs(text: str, n: int = 5) -> str:
    """Извлекает первые N абзацев из текста"""
    # Разбиваем текст на абзацы
    paragraphs = [p.strip() for p in text.split('\n') if p.strip()]
    
    # Берем первые n непустых абзацев
    selected = paragraphs[:n]
    return ' '.join(selected)


def clean_text(text: str) -> str:
    """Очистка текста от лишних символов"""
    # Убираем множественные пробелы
    text = re.sub(r'\s+', ' ', text)
    # Убираем неразрывные пробелы
    text = text.replace('\xa0', ' ')
    return text.strip()


def load_all_orders() -> pd.DataFrame:
    """
    Загружает все приказы из директории (TXT, DOCX, PDF)
    Возвращает DataFrame с колонками: file_name, full_text, text_for_embedding
    """
    orders = []
    
    if not os.path.exists(config.ORDERS_DIR):
        print(f"Директория {config.ORDERS_DIR} не найдена. Используйте processed_orders.csv")
        return pd.DataFrame(columns=['file_name', 'full_text', 'text_for_embedding'])
    
    files = [f for f in os.listdir(config.ORDERS_DIR) 
             if any(f.lower().endswith(ext) for ext in SUPPORTED_EXTENSIONS)]
    print(f"Найдено {len(files)} файлов приказов (TXT, DOCX, PDF)")
    
    for filename in tqdm(files, desc="Загрузка приказов"):
        file_path = os.path.join(config.ORDERS_DIR, filename)
        
        # Загружаем полный текст
        full_text = load_order_file(file_path)
        
        if len(full_text) < config.MIN_TEXT_LENGTH:
            continue
        
        # Извлекаем первые N абзацев для эмбеддингов
        text_for_embedding = extract_first_n_paragraphs(
            full_text, 
            config.MAX_PARAGRAPHS
        )
        text_for_embedding = clean_text(text_for_embedding)
        
        orders.append({
            'file_name': filename,
            'full_text': clean_text(full_text),
            'text_for_embedding': text_for_embedding
        })
    
    df = pd.DataFrame(orders)
    print(f"Успешно загружено {len(df)} приказов")
    
    return df


def load_error_patterns() -> Dict[str, List[Dict]]:
    """
    Загружает типовые ошибки и парсит их в структурированный формат
    Возвращает словарь с категориями ошибок и их описанием
    """
    error_categories = {}
    
    try:
        file_path = config.ERRORS_FILE
        
        # Пробуем разные кодировки
        encodings = ['utf-8', 'windows-1251', 'cp1251', 'utf-16']
        text = None
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    text = f.read()
                    if text and not text.startswith('�'):
                        break
            except (UnicodeDecodeError, UnicodeError):
                continue
        
        if not text:
            print("Не удалось корректно прочитать файл с ошибками")
            return create_default_error_patterns()
        
        # Парсим структуру ошибок
        # Формат: категория -> примеры -> правила
        current_category = None
        current_error = {}
        
        lines = text.split('\n')
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # Определяем новую категорию (начинается с цифры и точки)
            if re.match(r'^\d+\.', line):
                if current_category and current_error:
                    if current_category not in error_categories:
                        error_categories[current_category] = []
                    error_categories[current_category].append(current_error)
                
                current_category = line
                current_error = {
                    'description': line,
                    'examples': [],
                    'patterns': []
                }
            
            # Примеры ошибок (строки с "Пример:", "Ошибка:")
            elif any(keyword in line for keyword in ['Пример:', 'Ошибка:', 'Неверно:', 'Неправильно:']):
                if current_error:
                    current_error['examples'].append(line)
            
            # Правильные варианты
            elif any(keyword in line for keyword in ['Правильно:', 'Верно:', 'Следует:']):
                if current_error:
                    current_error['patterns'].append(line)
        
        # Добавляем последнюю категорию
        if current_category and current_error:
            if current_category not in error_categories:
                error_categories[current_category] = []
            error_categories[current_category].append(current_error)
        
        print(f"Загружено {len(error_categories)} категорий ошибок")
        
        return error_categories if error_categories else create_default_error_patterns()
        
    except Exception as e:
        print(f"Ошибка при загрузке файла ошибок: {e}")
        return create_default_error_patterns()


def create_default_error_patterns() -> Dict[str, List[Dict]]:
    """
    Создает базовые шаблоны ошибок на основе общих правил оформления приказов
    """
    return {
        "Форматирование дат": [{
            'description': 'Некорректное оформление дат',
            'examples': ['01 мая 2025', 'мая 2025 года'],
            'patterns': [
                r'\d{2}\s+[а-яё]+\s+\d{4}',  # неправильный формат даты
                r'[а-яё]+\s+\d{4}\s+года'     # месяц год года
            ]
        }],
        "Ссылки на документы": [{
            'description': 'Некорректные ссылки на нормативные документы',
            'examples': ['приказ без номера', 'постановление без даты'],
            'patterns': [
                r'приказ[^№\d]{10,}',  # приказ без номера
                r'постановление[^от]{10,}'  # постановление без даты
            ]
        }],
        "Пунктуация и пробелы": [{
            'description': 'Ошибки в пунктуации',
            'examples': ['точка после заголовка', 'двойные пробелы'],
            'patterns': [
                r'\s{2,}',  # множественные пробелы
                r'[А-ЯЁ][^\n]{20,}\.',  # точка после заголовка
            ]
        }],
        "Нумерация": [{
            'description': 'Некорректная нумерация пунктов',
            'examples': ['1.1., 1.2.', '1, 3, 2'],
            'patterns': [
                r'\d+\.\d+\.,',  # двойная точка
            ]
        }]
    }


def save_processed_data(df: pd.DataFrame, filename: str = "processed_orders.csv"):
    """Сохраняет обработанные данные"""
    output_path = os.path.join(config.OUTPUT_DIR, filename)
    df.to_csv(output_path, index=False, encoding='utf-8-sig')
    print(f"Данные сохранены в {output_path}")


if __name__ == "__main__":
    # Тестируем загрузку
    print("=== Загрузка приказов ===")
    df = load_all_orders()
    print(f"\nПример загруженного приказа:")
    print(f"Имя файла: {df.iloc[0]['file_name']}")
    print(f"Длина текста: {len(df.iloc[0]['full_text'])}")
    print(f"Текст для эмбеддинга (первые 200 символов):")
    print(df.iloc[0]['text_for_embedding'][:200])
    
    print("\n=== Загрузка типовых ошибок ===")
    errors = load_error_patterns()
    for category, patterns in list(errors.items())[:3]:
        print(f"\nКатегория: {category}")
        print(f"Количество паттернов: {len(patterns)}")
    
    # Сохраняем
    save_processed_data(df)
